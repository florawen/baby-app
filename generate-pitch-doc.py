from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Title
title = doc.add_heading('Baby Zhu-Van Swieten App', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('5-Minute Interview Pitch — AI Engineer Demo')
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()

# Opening
doc.add_heading('Opening (30 sec)', level=2)
doc.add_paragraph(
    '"I built this as a personal project for my sister\'s first baby — a family coordination app. '
    'What started as a simple idea became a real product I designed, architected, and shipped end-to-end '
    'using AI-assisted development with Claude. I want to walk you through not just what it does, but '
    'how I built it and the decisions I made."'
)

# The Problem
doc.add_heading('The Problem (30 sec)', level=2)
doc.add_paragraph(
    '"New parents are overwhelmed. They\'re tracking feeds, diapers, and sleep on scraps of paper. '
    'Their support network is scattered across contacts and texts. Family wants to help but doesn\'t '
    'know what\'s needed. I wanted one place where my sister\'s family could coordinate — and where '
    'guests could contribute without needing full access."'
)

# Architecture
doc.add_heading('Architecture & Tech Decisions (1 min)', level=2)
doc.add_paragraph(
    '"It\'s a static HTML/JS/CSS app — no framework, no build step. That was intentional. For a small '
    'family app, React or Vue would be overhead without benefit. I deployed to Firebase Hosting, which '
    'gives me HTTPS, CDN, and custom domain for free.'
)
doc.add_paragraph('For the backend, I used Firebase\'s serverless stack:')

bullets = [
    ('Authentication', 'with email/password — individual logins, not a shared PIN'),
    ('Firestore', 'for real-time data sync — when my sister logs a feed, everyone sees it instantly via snapshot listeners'),
    ('Firebase Storage', 'for photo uploads in the scrapbook'),
]
for bold_text, rest in bullets:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(bold_text)
    run.bold = True
    p.add_run(f' {rest}')

doc.add_paragraph(
    'The data model uses a family document with subcollections. Members and guests are tracked by UID, '
    'and role-based access controls what each user sees — guests only see the registry, family sees everything."'
)

# UI/UX
doc.add_heading('UI/UX Thinking (1 min)', level=2)
doc.add_paragraph(
    '"I designed for one-handed phone use — the primary user is holding a baby. Big tap targets, '
    'bottom navigation, quick-action buttons for the most common tasks (feed, diaper, sleep).'
)
doc.add_paragraph(
    'The app has seven tabs but I kept the nav tight with reduced padding and short labels. '
    'Each feature solves a real need:'
)

features = [
    ('Countdown', 'emotional, keeps the family excited'),
    ('Tracker', 'logs feeds, diapers, naps with timestamps and history'),
    ('Village', 'quick-dial contacts with role categories (pediatrician, midwife, family)'),
    ('Scrapbook', 'photo wall with lightbox viewer and navigation'),
    ('Journal', 'date-indexed entries with teasers'),
    ('Inventory', 'quantity tracking with low-stock alerts'),
    ('Registry', 'guest-accessible wishlist with claiming'),
]
for bold_text, rest in features:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(bold_text)
    run.bold = True
    p.add_run(f' — {rest}')

doc.add_paragraph(
    'The guest experience was a specific design challenge: guests register through invite links, '
    'see only the registry, can claim items, and other guests see \'someone\'s getting this\' without '
    'names — preventing duplicates while preserving surprise."'
)

# AI-Assisted Development
doc.add_heading('AI-Assisted Development Process (1 min)', level=2)
doc.add_paragraph(
    '"I used Claude Code as my development partner throughout. The workflow was conversational — '
    'I\'d describe what I wanted in plain language, we\'d discuss trade-offs (one tab vs two for '
    'inventory/registry, PIN auth vs individual logins, toggle views vs separate tabs), and then '
    'Claude would implement while I reviewed and tested.'
)

doc.add_paragraph('Key things I directed:')
directed = [
    'The architectural decisions (Firebase over alternatives, no framework, role-based access model)',
    'The UX priorities (mobile-first, one-handed use, progressive disclosure)',
    'Feature scoping (what\'s MVP vs. what\'s future)',
]
for item in directed:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph('What AI accelerated:')
accelerated = [
    'Boilerplate (Firebase config, CRUD operations, CSS grid layouts)',
    'Error handling patterns (offline persistence, listener cleanup, try/catch isolation)',
    'Iterative refinement (lightbox navigation, countdown timer, demo mode)',
]
for item in accelerated:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph(
    'This is how I see AI engineering working — I\'m the product thinker and architect, AI handles '
    'velocity on implementation, and we iterate together."'
)

# Resilience
doc.add_heading('Resilience & Production Thinking (30 sec)', level=2)
doc.add_paragraph(
    '"I built for real-world reliability: Firestore offline persistence so the app works without signal, '
    'snapshot error callbacks so listeners don\'t die silently, module initialization wrapped in try/catch '
    'so one broken feature doesn\'t take down the whole app, and a service worker for PWA-like behavior. '
    'Storage rules enforce that only authenticated users in the family can read/write data."'
)

# Close
doc.add_heading('Close (30 sec)', level=2)
doc.add_paragraph(
    '"This is a shipped product — my family uses it daily. It demonstrates end-to-end product thinking: '
    'identifying a user need, making architectural trade-offs, designing for accessibility and edge cases, '
    'and leveraging AI tools to ship faster without sacrificing quality. The demo mode you\'re looking at '
    'loads entirely client-side with fake data — the same codebase, zero personal information exposed."'
)

# Demo link
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Demo: baby-zhu-van-swieten.web.app?demo')
run.bold = True
run.font.size = Pt(12)

output_path = '/Users/FLORZHU@tbdir.net/baby-app/Baby-App-Interview-Pitch.docx'
doc.save(output_path)
print(f'Saved to {output_path}')
